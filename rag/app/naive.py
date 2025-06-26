#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
from functools import reduce
from io import BytesIO
from timeit import default_timer as timer

from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from markdown import markdown
from PIL import Image
from tika import parser

from api.db import LLMType
from api.db.services.llm_service import LLMBundle
from api.utils.langfuse_utils import LangfuseUtils, create_span_with_langfuse
from deepdoc.parser import DocxParser, ExcelParser, HtmlParser, JsonParser, MarkdownParser, PdfParser, TxtParser
from deepdoc.parser.figure_parser import VisionFigureParser, vision_figure_parser_figure_data_wraper, pdf_vision_figure_parser_figure_data_wraper
from deepdoc.parser.pdf_parser import PlainParser, VisionParser
from rag.nlp import concat_img, find_codec, naive_merge, naive_merge_with_images, naive_merge_docx, rag_tokenizer, tokenize_chunks, tokenize_chunks_with_images, tokenize_table
from rag.utils import num_tokens_from_string


def _trace_sections_and_tables_to_langfuse(sections, tables, main_trace):
    """在主 trace 下添加 sections 和 tables 内容的 span"""
    if not main_trace:
        return
    
    try:
        # 处理 sections 内容
        sections_data = []
        for i, section in enumerate(sections):
            section_info = {
                "index": i,
                "text": section[0][:500] if section[0] else "",  # 前500字符
                "has_image": section[1] is not None,
                "style": section[2] if len(section) > 2 else "unknown",
                "text_length": len(section[0]) if section[0] else 0
            }
            
            # 如果有图片，尝试编码为base64（用于展示）
            if section[1] is not None:
                try:
                    import base64
                    from io import BytesIO
                    img_buffer = BytesIO()
                    section[1].save(img_buffer, format='PNG')
                    img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                    section_info["image_preview"] = f"data:image/png;base64,{img_base64}"
                except Exception:
                    section_info["image_preview"] = "Failed to encode image"
            
            sections_data.append(section_info)
        
        # 处理 tables 内容
        tables_data = []
        for i, table in enumerate(tables):
            table_info = {
                "index": i,
                "html_content": table[0][1][:1000] if table[0] and table[0][1] else "",  # 前1000字符
                "content_length": len(table[0][1]) if table[0] and table[0][1] else 0,
                "table_type": table[1] if len(table) > 1 else "unknown"
            }
            tables_data.append(table_info)
        
        # 创建 sections 和 tables 的联合 span
        content_span = create_span_with_langfuse(
            main_trace,
            name="sections_and_tables",
            input_data={
                "sections_count": len(sections),
                "tables_count": len(tables)
            },
            output_data={
                "sections_preview": sections_data,
                "tables_preview": tables_data,
                "summary": {
                    "total_sections": len(sections),
                    "total_tables": len(tables),
                    "sections_with_images": sum(1 for s in sections if s[1] is not None)
                }
            }
        )
        
    except Exception as e:
        logging.warning(f"Failed to trace sections and tables to Langfuse: {e}")


def _trace_figure_enhancement_to_langfuse(original_figures_data, enhanced_figures_data, main_trace):
    """在主 trace 下添加图像增强结果的 span，同时包含原始和增强后的信息"""
    if not main_trace or not enhanced_figures_data:
        return
    
    try:
        # 创建原始图片数据的索引，用于匹配增强后的图片
        original_images_info = []
        for i, original_figure in enumerate(original_figures_data):
            if isinstance(original_figure, tuple) and len(original_figure) >= 1:
                original_content = original_figure[0]
                if isinstance(original_content, tuple) and len(original_content) >= 2:
                    image_data, original_descriptions = original_content[0], original_content[1]
                    if isinstance(original_descriptions, list) and original_descriptions:
                        # 这里就是包含"Caption: ... ||| Context: ..."的原始描述
                        original_context = original_descriptions[0] if original_descriptions[0] else ""
                        original_images_info.append({
                            "original_index": i,
                            "image_data": image_data,
                            "original_context": original_context
                        })
        
        # 处理增强后的图像数据
        enhanced_figures = []
        for i, enhanced_figure in enumerate(enhanced_figures_data[:5]):  # 只取前5个
            figure_info = {
                "enhanced_index": i,
                "enhanced_description": "",
                "generated_image_name": "",
                "original_context": "",
                "original_index": -1,
                "has_image": False
            }
            
            # 从增强后的数据中提取信息
            if isinstance(enhanced_figure, tuple) and len(enhanced_figure) >= 2:
                # 格式: ((image, descriptions, filename), positions)
                content, positions = enhanced_figure[0], enhanced_figure[1] if len(enhanced_figure) > 1 else []
                if isinstance(content, tuple) and len(content) >= 2:
                    image_data = content[0]
                    descriptions = content[1]
                    filename = content[2] if len(content) >= 3 else ""
                    
                    # 提取增强后的描述（VisionFigureParser生成的）
                    if isinstance(descriptions, list) and descriptions:
                        enhanced_desc = descriptions[0] if descriptions[0] else ""
                        figure_info["enhanced_description"] = enhanced_desc[:500] if enhanced_desc else ""
                    
                    # 提取大模型生成的image_name
                    figure_info["generated_image_name"] = filename if filename else ""
                    
                    # 通过图片数据匹配原始上下文信息
                    # 由于VisionFigureParser可能会跳过一些图片，我们需要通过图片内容来匹配
                    if image_data is not None:
                        figure_info["has_image"] = True
                        
                        # 尝试匹配原始图片信息
                        for orig_info in original_images_info:
                            if orig_info["image_data"] is image_data:  # 同一个图片对象
                                figure_info["original_context"] = orig_info["original_context"][:500] if orig_info["original_context"] else ""
                                figure_info["original_index"] = orig_info["original_index"]
                                break
                        
                        # 编码图片预览
                        try:
                            import base64
                            from io import BytesIO
                            img_buffer = BytesIO()
                            image_data.save(img_buffer, format='PNG')
                            img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                            figure_info["image_preview"] = f"data:image/png;base64,{img_base64}"
                        except Exception:
                            figure_info["image_preview"] = "Failed to encode enhanced image"
            
            enhanced_figures.append(figure_info)
        
        # 在主 trace 下创建图像增强 span
        enhancement_span = create_span_with_langfuse(
            main_trace,
            name="figure_enhancement",
            input_data={
                "original_figures_count": len(original_figures_data),
                "enhanced_figures_count": len(enhanced_figures_data)
            },
            output_data={
                "enhanced_figures": enhanced_figures,
                "summary": {
                    "total_original": len(original_figures_data),
                    "total_enhanced": len(enhanced_figures_data),
                    "figures_with_descriptions": sum(1 for f in enhanced_figures if f["enhanced_description"]),
                    "figures_with_generated_names": sum(1 for f in enhanced_figures if f["generated_image_name"]),
                    "figures_with_images": sum(1 for f in enhanced_figures if f["has_image"]),
                    "figures_with_original_context": sum(1 for f in enhanced_figures if f["original_context"]),
                    "figures_matched_to_original": sum(1 for f in enhanced_figures if f["original_index"] >= 0)
                }
            }
        )
        
        # 注意：主 trace 的最终更新在主流程中完成
        
    except Exception as e:
        logging.warning(f"Failed to trace figure enhancement to Langfuse: {e}")


def _trace_pdf_sections_tables_figures_to_langfuse(sections, tables, figures, main_trace):
    """在主 trace 下添加 PDF sections、tables 和 figures 的 span"""
    if not main_trace:
        return
    
    try:
        # 处理 sections 内容
        sections_data = []
        for i, section in enumerate(sections):
            section_info = {
                "index": i,
                "content_preview": "",
                "line_tag": "",
                "content_length": 0
            }
            
            if isinstance(section, tuple) and len(section) >= 2:
                content, line_tag = section[0], section[1]
                section_info["content_preview"] = content[:200] if content else ""
                section_info["line_tag"] = line_tag[:100] if line_tag else ""
                section_info["content_length"] = len(content) if content else 0
            elif isinstance(section, str):
                section_info["content_preview"] = section[:200]
                section_info["content_length"] = len(section)
            
            sections_data.append(section_info)
        
        # 处理 tables 内容
        tables_data = []
        for i, table in enumerate(tables):
            table_info = {
                "index": i,
                "table_preview": "",
                "table_type": "",
                "content_length": 0
            }
            
            # PDF table 数据结构可能比较复杂，这里做简单处理
            if isinstance(table, tuple) and len(table) >= 2:
                table_content = table[0]
                if isinstance(table_content, str):
                    table_info["table_preview"] = table_content[:300]
                    table_info["content_length"] = len(table_content)
                elif hasattr(table_content, '__str__'):
                    content_str = str(table_content)
                    table_info["table_preview"] = content_str[:300]
                    table_info["content_length"] = len(content_str)
                table_info["table_type"] = "pdf_extracted_table"
            elif isinstance(table, str):
                table_info["table_preview"] = table[:300]
                table_info["content_length"] = len(table)
                table_info["table_type"] = "text_table"
            else:
                table_info["table_type"] = str(type(table))
            
            tables_data.append(table_info)
        
        # 处理 figures 内容
        figures_data = []
        for i, figure in enumerate(figures):
            figure_info = {
                "index": i,
                "has_image": False,
                "captions": [],
                "positions": [],
                "figure_type": ""
            }
            
            # PDF figure 数据结构: ((image, captions), positions)
            if isinstance(figure, tuple) and len(figure) >= 2:
                img_desc, positions = figure[0], figure[1]
                if isinstance(img_desc, tuple) and len(img_desc) >= 2:
                    image_data, captions = img_desc[0], img_desc[1]
                    
                    figure_info["has_image"] = image_data is not None
                    figure_info["captions"] = captions[:3] if isinstance(captions, list) else []  # 只取前3个caption
                    figure_info["positions"] = positions[:3] if isinstance(positions, list) else []  # 只取前3个位置
                    figure_info["figure_type"] = "pdf_extracted_figure"
                    
                    # 如果有图片，尝试编码预览
                    if image_data is not None:
                        try:
                            import base64
                            from io import BytesIO
                            img_buffer = BytesIO()
                            image_data.save(img_buffer, format='PNG')
                            img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                            figure_info["image_preview"] = f"data:image/png;base64,{img_base64}"
                        except Exception:
                            figure_info["image_preview"] = "Failed to encode figure image"
            else:
                figure_info["figure_type"] = str(type(figure))
                
            figures_data.append(figure_info)
        
        # 在主 trace 下创建 PDF 内容 span
        content_span = create_span_with_langfuse(
            main_trace,
            name="pdf_content_extraction",
            input_data={"document_type": "pdf"},
            output_data={
                "sections": sections_data,
                "tables": tables_data,
                "figures": figures_data,
                "summary": {
                    "total_sections": len(sections),
                    "total_tables": len(tables),
                    "total_figures": len(figures),
                    "sections_with_content": sum(1 for s in sections_data if s["content_length"] > 0),
                    "tables_with_content": sum(1 for t in tables_data if t["content_length"] > 0),
                    "figures_with_images": sum(1 for f in figures_data if f["has_image"]),
                    "figures_with_captions": sum(1 for f in figures_data if f["captions"])
                }
            }
        )
        
    except Exception as e:
        logging.warning(f"Failed to trace PDF sections, tables and figures to Langfuse: {e}")


def _trace_chunks_merging_to_langfuse(sections, chunks, images, parser_config, main_trace):
    """在主 trace 下添加 sections 合并成 chunks 的 span"""
    if not main_trace:
        return
    
    try:
        # 收集合并前的信息
        original_sections_info = []
        for i, section in enumerate(sections):  # 只取前10个sections进行详细记录
            section_info = {
                "index": i,
                "content_preview": "",
                "content_length": 0,
                "has_line_tag": False
            }
            
            if isinstance(section, tuple) and len(section) >= 2:
                content, line_tag = section[0], section[1]
                section_info["content_preview"] = content[:200] if content else ""
                section_info["content_length"] = len(content) if content else 0
                section_info["has_line_tag"] = bool(line_tag)
            elif isinstance(section, str):
                section_info["content_preview"] = section[:200]
                section_info["content_length"] = len(section)
            
            original_sections_info.append(section_info)
        
        # 收集合并后的chunks信息
        chunks_info = []
        for i, chunk in enumerate(chunks):  # 只取前10个chunks进行详细记录
            chunk_info = {
                "index": i,
                "content_preview": chunk if chunk else "",
                "content_length": len(chunk) if chunk else 0
            }
            chunks_info.append(chunk_info)
        
        # 收集images信息
        images_info = []
        if images:
            for i, image in enumerate(images):  # 只取前5个images
                image_info = {
                    "index": i,
                    "has_image": image is not None
                }
                
                # 如果有图片，尝试编码预览
                if image is not None:
                    try:
                        import base64
                        from io import BytesIO
                        img_buffer = BytesIO()
                        image.save(img_buffer, format='PNG')
                        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                        image_info["image_preview"] = f"data:image/png;base64,{img_base64}"
                    except Exception:
                        image_info["image_preview"] = "Failed to encode chunk image"
                
                images_info.append(image_info)
        
        # 在主 trace 下创建 chunks 合并 span
        merging_span = create_span_with_langfuse(
            main_trace,
            name="sections_to_chunks_merging",
            input_data={
                "chunk_token_num": parser_config.get("chunk_token_num", 128),
                "delimiter": parser_config.get("delimiter", "\n!?。；！？"),
                "original_sections_count": len(sections)
            },
            output_data={
                "original_sections": original_sections_info,
                "merged_chunks": chunks_info,
                "chunk_images": images_info,
                "summary": {
                    "original_sections_count": len(sections),
                    "merged_chunks_count": len(chunks),
                    "chunks_with_images": len(images) if images else 0,
                    "sections_with_content": sum(1 for s in original_sections_info if s["content_length"] > 0),
                    "average_section_length": sum(s["content_length"] for s in original_sections_info) / len(original_sections_info) if original_sections_info else 0,
                    "average_chunk_length": sum(c["content_length"] for c in chunks_info) / len(chunks_info) if chunks_info else 0,
                    "compression_ratio": len(sections) / len(chunks) if chunks else 0
                }
            }
        )
        
    except Exception as e:
        logging.warning(f"Failed to trace chunks merging to Langfuse: {e}")


# 示例：如何添加新的追踪步骤
def _trace_tokenization_to_langfuse(chunks, main_trace):
    """示例：追踪 tokenization 步骤"""
    if not main_trace:
        return
    
    try:
        tokenization_span = create_span_with_langfuse(
            main_trace,
            name="tokenization",
            input_data={"chunks_count": len(chunks)},
            output_data={
                "summary": {
                    "total_chunks": len(chunks),
                    "chunks_preview": [chunk[:100] for chunk in chunks[:3]]  # 前3个chunk的预览
                }
            }
        )
    except Exception as e:
        logging.warning(f"Failed to trace tokenization to Langfuse: {e}")


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')
        if not embed:
            return None
        embed = embed[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        except UnicodeDecodeError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        # 替换全角空格为半角空格
        line = re.sub(r"\u3000", " ", line)
        
        # 使用更全面的正则表达式处理所有中文字符之间的空格
        # 匹配任意数量的中文字符，中间可能有空格
        line = re.sub(r'([\u4e00-\u9fff]+)\s+([\u4e00-\u9fff]+)', r'\1\2', line)
        
        # 去除中文字符与数字之间的空格
        line = re.sub(r'([\u4e00-\u9fff])\s+(\d)', r'\1\2', line)
        line = re.sub(r'(\d)\s+([\u4e00-\u9fff])', r'\1\2', line)
        
        # 去除中文字符与英文之间的空格
        line = re.sub(r'([\u4e00-\u9fff])\s+([a-zA-Z])', r'\1\2', line)
        line = re.sub(r'([a-zA-Z])\s+([\u4e00-\u9fff])', r'\1\2', line)
        
        # 去除中文字符与标点符号之间的空格
        # 转义引号并添加更多标点符号
        line = re.sub(r'([\u4e00-\u9fff])\s+([，。！？；：\"\'（）【】《》、/\\])', r'\1\2', line)
        line = re.sub(r'([，。！？；：\"\'（）【】《》、/\\])\s+([\u4e00-\u9fff])', r'\2\1', line)
        
        # 去除重复的空格
        line = re.sub(r'\s+', ' ', line)
        
        # 去除首尾空格
        line = line.strip()
        
        # 再次检查是否还有中文字符之间的空格（处理可能遗漏的情况）
        line = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', line)
        
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph
        
        titles = []
        blocks = []
        
        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"
            
        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith('p'):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(('p', i, p))
                elif block.tag.endswith('tbl'):  # Table
                    blocks.append(('t', i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""
            
        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == 't':
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1
                
        if target_table_pos == -1:
            return ""  # Target table not found
            
        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks)-1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue
                
            if block_type != 'p':
                continue
                
            if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")
        
        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]
            
            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks)-1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue
                        
                    if block_type != 'p':
                        continue
                        
                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:  
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")
                            
                if not found:  # Break if no parent heading is found
                    break
            
            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)
            
        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for i, tb in enumerate(self.doc.tables):
            title = self.__get_nearest_title(i, filename)
            html = "<table>"
            if title:
                html += f"<caption>Table Location: {title}</caption>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        if separate_tables_figures:
            # 调用提取函数
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            # self._naive_vertical_merge()
            self._concat_downward()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def get_picture_urls(self, sections):
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []
        
        from bs4 import BeautifulSoup
        html_content = markdown(text)
        soup = BeautifulSoup(html_content, 'html.parser')
        html_images = [img.get('src') for img in soup.find_all('img') if img.get('src')]
        return html_images


    def get_pictures(self, text):
        """Download and open all images from markdown text."""
        import requests
        image_urls = self.get_picture_urls(text)
        images = []
        # Find all image URLs in text
        for url in image_urls:
            try:
                response = requests.get(url, stream=True, timeout=30)
                if response.status_code == 200 and response.headers['Content-Type'].startswith('image/'):
                    img = Image.open(BytesIO(response.content)).convert('RGB')
                    images.append(img)
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
                continue

        return images if images else None

    def __call__(self, filename, binary=None):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        sections = []
        tbls = []
        for sec in remainder.split("\n"):
            if num_tokens_from_string(sec) > 3 * self.chunk_token_num:
                sections.append((sec[:int(len(sec) / 2)], ""))
                sections.append((sec[int(len(sec) / 2):], ""))
            else:
                if sec.strip().find("#") == 0:
                    sections.append((sec, ""))
                elif sections and sections[-1][0].strip().find("#") == 0:
                    sec_, _ = sections.pop(-1)
                    sections.append((sec_ + "\n" + sec, ""))
                else:
                    sections.append((sec, ""))
        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, excel, txt.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """
    
    # 记录全过程开始时间
    start_time = timer()
    
    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    section_images = None
    main_trace = None  # 初始化主trace变量，用于所有文档类型
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 在主流程开始时创建 Langfuse 主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="docx_document_analysis",
                input_data={"filename": filename},
                metadata={"type": "docx_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for DOCX: {e}")

        try:
            vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT)
            callback(0.15, "Visual model detected. Attempting to enhance figure extraction...")
        except Exception as e:
            logging.error(f"Failed to create vision model: {e}")
            vision_model = None

        sections, tables = Docx()(filename, binary)
        
        # 追踪 sections 和 tables 内容
        _trace_sections_and_tables_to_langfuse(sections, tables, main_trace)
        
        figures_data = []
        if vision_model:
            logging.info(f'vision_model is {vision_model}, start to enhance figure extraction...')
            try:
                # 使用增强的包装函数处理图片数据
                original_figures_data = vision_figure_parser_figure_data_wraper(sections, context_window=5)
                
                docx_vision_parser = VisionFigureParser(
                    vision_model=vision_model, 
                    figures_data=original_figures_data,
                    **kwargs
                )
                boosted_figures = docx_vision_parser(callback=callback)
                figures_data = boosted_figures
                
                # 追踪图像增强结果，传递原始和增强后的数据
                _trace_figure_enhancement_to_langfuse(original_figures_data, figures_data, main_trace)
                
            except Exception as e:
                callback(0.6, f"Visual model error: {e}. Skipping figure parsing enhancement.")

        res = tokenize_table(tables, doc, is_english)
        if figures_data:
            res.extend(tokenize_table(figures_data, doc, is_english, table_type="figure"))
        callback(0.8, "Finish parsing.")

        # 这里可以很容易地添加其他步骤的追踪，例如：
        # _trace_tokenization_to_langfuse(res, main_trace)

        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        # 追踪 sections 合并成 chunks 的过程
        _trace_chunks_merging_to_langfuse(sections, chunks, images, parser_config, main_trace)

        if kwargs.get("section_only", False):
            final_result = tokenize_chunks_with_images(chunks, doc, is_english, images)
            # 更新主 trace 的最终结果
            if main_trace:
                LangfuseUtils.update_trace(main_trace, {
                    "section_only": True,
                    "final_chunks_count": len(final_result),
                    "analysis_complete": True,
                    "processing_time_seconds": timer() - start_time
                })
            return final_result

        res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
        
        # 更新主 trace 的最终结果
        if main_trace:
            LangfuseUtils.update_trace(main_trace, {
                "final_chunks_count": len(res),
                "analysis_complete": True,
                "processing_time_seconds": timer() - start_time
            })
        
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"
        callback(0.1, "Start to parse.")

        # 为PDF分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="pdf_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "pdf_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for PDF: {e}")

        if layout_recognizer == "DeepDOC":
            pdf_parser = Pdf()

            try:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT)
                callback(0.15, "Visual model detected. Attempting to enhance figure extraction...")
            except Exception:
                vision_model = None

            sections, tables, figures = pdf_parser(filename if not binary else binary, 
                                                   from_page=from_page, 
                                                   to_page=to_page, 
                                                   callback=callback, 
                                                   separate_tables_figures=True, 
                                                   zoomin=5)
            
            # 追踪 PDF sections、tables 和 figures 内容
            _trace_pdf_sections_tables_figures_to_langfuse(sections, tables, figures, main_trace)
            
            if vision_model:
                callback(0.5, "Basic parsing complete. Proceeding with figure enhancement...")
                try:
                    # 使用新的PDF图片上下文包装函数
                    original_figures_data = pdf_vision_figure_parser_figure_data_wraper(
                        figures, sections, min_context_chars=kwargs.get("min_context_chars", 800)
                    )
                    pdf_vision_parser = VisionFigureParser(
                        vision_model=vision_model, 
                        figures_data=original_figures_data, 
                        **kwargs
                    )
                    boosted_figures = pdf_vision_parser(callback=callback)
                    # 将增强后的图片数据单独保留
                    figures = boosted_figures
                    
                    # 追踪PDF图像增强结果，传递原始和增强后的数据
                    _trace_figure_enhancement_to_langfuse(original_figures_data, figures, main_trace)
                    
                except Exception as e:
                    logging.error(f"Vision model error during figure enhancement: {e}")
                    callback(0.6, f"Visual model error: {e}. Skipping figure parsing enhancement.")
            # 分别处理tables和figures
            res = tokenize_table(tables, doc, is_english)
            if figures:
                res.extend(tokenize_table(figures, doc, is_english, table_type="figure"))
            
            callback(0.8, "Finish parsing.")

        else:
            if layout_recognizer == "Plain Text":
                pdf_parser = PlainParser()
            else:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT, llm_name=layout_recognizer, lang=lang)
                pdf_parser = VisionParser(vision_model=vision_model, **kwargs)

            sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page,
                                          callback=callback)
            
            # 追踪 PDF sections 和 tables 内容（非DeepDOC模式）
            _trace_sections_and_tables_to_langfuse(sections, tables, main_trace)
            
            res = tokenize_table(tables, doc, is_english)
            callback(0.8, "Finish parsing.")

    elif re.search(r"\.(csv|xlsx?)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为Excel分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="excel_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "excel_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for Excel: {e}")
        
        excel_parser = ExcelParser()
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为文本文件分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="text_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "text_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for text file: {e}")
        
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为Markdown分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="markdown_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "markdown_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for Markdown: {e}")
        
        markdown_parser = Markdown(int(parser_config.get("chunk_token_num", 128)))
        sections, tables = markdown_parser(filename, binary)

        # 追踪 Markdown sections 和 tables 内容
        _trace_sections_and_tables_to_langfuse(sections, tables, main_trace)

        # Process images for each section
        section_images = []
        for section_text, _ in sections:
            images = markdown_parser.get_pictures(section_text) if section_text else None
            if images:
                # If multiple images found, combine them using concat_img
                combined_image = reduce(concat_img, images) if len(images) > 1 else images[0]
                section_images.append(combined_image)
            else:
                section_images.append(None)

        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为HTML分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="html_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "html_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for HTML: {e}")
        
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为JSON分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="json_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "json_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for JSON: {e}")
        
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        
        # 为DOC分析创建主 trace
        try:
            langfuse_client = LangfuseUtils.get_client(kwargs.get("tenant_id"))
            main_trace = LangfuseUtils.create_trace(
                langfuse_client,
                name="doc_document_analysis", 
                input_data={"filename": filename},
                metadata={"type": "doc_analysis", "document": filename}
            )
        except Exception as e:
            logging.warning(f"Failed to create main trace for DOC: {e}")
        
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    if section_images:
        # if all images are None, set section_images to None
        if all(image is None for image in section_images):
            section_images = None

    if section_images:
        chunks, images = naive_merge_with_images(sections, section_images,
                                        int(parser_config.get(
                                            "chunk_token_num", 128)), parser_config.get(
                                            "delimiter", "\n!?。；！？"))
        
        # 追踪 sections 合并成 chunks 的过程（带图片）
        _trace_chunks_merging_to_langfuse(sections, chunks, images, parser_config, main_trace)
        
        if kwargs.get("section_only", False):
            final_result = tokenize_chunks(chunks, doc, is_english, pdf_parser)
            # 更新主 trace 的最终结果（对于PDF等其他文件类型）
            if main_trace:
                LangfuseUtils.update_trace(main_trace, {
                    "section_only": True,
                    "final_chunks_count": len(final_result),
                    "analysis_complete": True,
                    "processing_time_seconds": timer() - start_time
                })
            return final_result

        res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
    else:
        chunks = naive_merge(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))
        
        # 追踪 sections 合并成 chunks 的过程（无图片）
        _trace_chunks_merging_to_langfuse(sections, chunks, None, parser_config, main_trace)
        
        if kwargs.get("section_only", False):
            final_result = tokenize_chunks(chunks, doc, is_english, pdf_parser)
            # 更新主 trace 的最终结果（对于PDF等其他文件类型）
            if main_trace:
                LangfuseUtils.update_trace(main_trace, {
                    "section_only": True,
                    "final_chunks_count": len(final_result),
                    "analysis_complete": True,
                    "processing_time_seconds": timer() - start_time
                })
            return final_result

        res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))

    # 更新主 trace 的最终结果（对于PDF等其他文件类型）
    if main_trace:
        LangfuseUtils.update_trace(main_trace, {
            "final_chunks_count": len(res),
            "analysis_complete": True,
            "processing_time_seconds": timer() - start_time
        })

    logging.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
